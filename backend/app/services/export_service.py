"""
Export service for different file formats
"""
import pandas as pd
from docx import Document
from typing import List
import io
from fastapi.responses import StreamingResponse
from models.schemas import SequenceItem

class ExportService:
    async def export(self, sequence: List[SequenceItem], format: str):
        if format == "csv":
            return await self._export_csv(sequence)
        elif format == "docx":
            return await self._export_docx(sequence)
        elif format == "txt":
            return await self._export_txt(sequence)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    async def _export_csv(self, sequence: List[SequenceItem]):
        """Export to CSV format"""
        data = []
        for item in sequence:
            data.append({
                "Step": item.step,
                "Purpose": item.purpose,
                "Subject": item.subject,
                "Body": item.body
            })
        
        df = pd.DataFrame(data)
        
        stream = io.StringIO()
        df.to_csv(stream, index=False)
        stream.seek(0)
        
        return StreamingResponse(
            iter([stream.getvalue()]),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename=sequence.csv"}
        )
    
    async def _export_docx(self, sequence: List[SequenceItem]):
        """Export to DOCX format"""
        doc = Document()
        doc.add_heading('Content Sequence', 0)
        
        for item in sequence:
            doc.add_heading(f'Step {item.step}: {item.purpose}', level=1)
            doc.add_heading(f'Subject: {item.subject}', level=2)
            doc.add_paragraph(item.body)
            doc.add_paragraph()  # Empty line
        
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        
        return StreamingResponse(
            iter([stream.getvalue()]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=sequence.docx"}
        )
    
    async def _export_txt(self, sequence: List[SequenceItem]):
        """Export to TXT format"""
        content = "Content Sequence\n================\n\n"
        
        for item in sequence:
            content += f"Step {item.step}: {item.purpose}\n"
            content += f"Subject: {item.subject}\n"
            content += f"Body: {item.body}\n\n"
            content += "-" * 50 + "\n\n"
        
        return StreamingResponse(
            iter([content]),
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=sequence.txt"}
        )